"""Unit tests for ingest/classifier.py — auto-tagging heuristics."""

from pathlib import Path

import pytest

from ingest.classifier import (
    DOC_TYPE_RULES,
    PRODUCT_ALIASES,
    PRODUCT_FROM_NAME,
    _build_metadata_text,
    classify,
    classify_document,
    enrich_classification,
    extract_document_metadata,
    infer_doc_type,
    infer_product,
    infer_vendor,
    infer_subsystem,
)


class TestInferDocType:
    def test_admin_guide(self):
        assert (
            infer_doc_type(Path("Content-Server-Admin-Guide.pdf"))
            == "admin_guide"
        )

    def test_standard_iso(self):
        assert infer_doc_type(Path("ISO-9001-Standard.pdf")) == "standard"

    def test_training(self):
        assert infer_doc_type(Path("Training-Manual-v2.pdf")) == "training"

    def test_release_notes(self):
        assert (
            infer_doc_type(Path("Release-Notes-23.4.pdf")) == "release_notes"
        )

    def test_install_guide(self):
        assert (
            infer_doc_type(Path("Installation-Guide.pdf")) == "install_guide"
        )

    def test_user_guide(self):
        assert infer_doc_type(Path("User-Guide.pdf")) == "user_guide"

    def test_api_guide(self):
        assert infer_doc_type(Path("API-Reference.pdf")) == "api_guide"

    def test_howto(self):
        assert infer_doc_type(Path("Troubleshooting-Guide.pdf")) == "howto"

    def test_reference(self):
        assert infer_doc_type(Path("Technical-Reference.pdf")) == "reference"

    def test_file_without_recognized_doc_type(self):
        assert infer_doc_type(Path("README.txt")) == "document"

    def test_zip_is_release_artifact(self):
        assert infer_doc_type(Path("some-package.zip")) == "release_artifact"

    def test_empty_path_does_not_crash(self):
        result = infer_doc_type(Path(""))
        assert isinstance(result, str)

    def test_case_insensitive_matching(self):
        assert infer_doc_type(Path("ADMIN-GUIDE.PDF")) == "admin_guide"


class TestInferProduct:
    def test_product_from_directory_name(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "ContentServer" / "admin.pdf"
        docs_root.mkdir()
        (docs_root / "ContentServer").mkdir()
        (docs_root / "ContentServer" / "admin.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "ContentServer"

    def test_product_from_override(self, tmp_path):
        file_path = tmp_path / "file.pdf"
        assert (
            infer_product(file_path, tmp_path, product_override="xECM")
            == "xECM"
        )

    def test_product_from_alias(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "appserver" / "admin.pdf"
        docs_root.mkdir()
        (docs_root / "appserver").mkdir()
        (docs_root / "appserver" / "admin.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "AppServer"

    def test_product_from_filename(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "data-sync-config.pdf"
        docs_root.mkdir()
        (docs_root / "data-sync-config.pdf").write_text("")
        product = infer_product(file_path, docs_root)
        assert product == "DataSync"

    def test_product_fallback_geral(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "varios" / "misc.pdf"
        docs_root.mkdir()
        (docs_root / "varios").mkdir()
        (docs_root / "varios" / "misc.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "geral"

    def test_product_not_in_alias_uses_dir_name(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "CustomProduct" / "file.pdf"
        docs_root.mkdir()
        (docs_root / "CustomProduct").mkdir()
        (docs_root / "CustomProduct" / "file.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "CustomProduct"


class TestClassify:
    def test_classify_basic(self, tmp_path):
        docs_root = tmp_path / "docs"
        product_dir = docs_root / "ContentServer"
        product_dir.mkdir(parents=True)
        file_path = product_dir / "Admin-Guide-23.4.pdf"
        file_path.write_text("")
        result = classify(file_path, docs_root)
        assert result["product"] == "ContentServer"
        assert result["doc_type"] == "admin_guide"
        # New keys present (Phase 11)
        assert "vendor" in result
        assert "subsystem" in result

    def test_classify_with_override(self, tmp_path):
        file_path = tmp_path / "file.pdf"
        result = classify(file_path, tmp_path, product_override="xECM")
        assert result["product"] == "xECM"

    def test_classify_document_wrapper(self, tmp_path):
        file_path = tmp_path / "Training-Material.pdf"
        result = classify_document(file_path)
        assert result["doc_type"] == "training"

    def test_classify_no_crash_with_unknown_file(self, tmp_path):
        file_path = tmp_path / "random_notes.txt"
        result = classify(file_path, tmp_path)
        assert isinstance(result, dict)
        assert "product" in result and "doc_type" in result

    def test_classify_returns_vendor_and_subsystem(self, tmp_path):
        """Verify classify() returns vendor and subsystem keys."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "WebReports" / "API-Guide.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_text("")
        result = classify(file_path, docs_root)
        assert "vendor" in result
        assert "subsystem" in result
        assert result["vendor"] == "OpenText"  # via VENDOR_MAP for WebReports
        assert result["subsystem"] == "API"  # via filename pattern

    def test_classify_opentext_webreport_admin_guide(self, tmp_path):
        """SC1: 'OpenText WebReport Administrator Guide 23.4.pdf' → all fields correct.

        This is the success criteria test from 11-01-PLAN.md SC1.
        """
        docs_root = tmp_path / "docs"
        # Using classify_document which uses parent as docs_root
        file_path = (
            docs_root / "OpenText WebReport Administrator Guide 23.4.pdf"
        )
        docs_root.mkdir()
        file_path.write_text("")
        result = classify_document(file_path)
        assert result["vendor"] == "OpenText"
        assert result["product"] == "WebReports"
        assert result["doc_type"] == "admin_guide"
        assert result["version"] == "23.4"


class TestModuleConstants:
    def test_doc_type_rules_are_well_formed(self):
        for priority, doc_type, patterns in DOC_TYPE_RULES:
            assert isinstance(priority, int)
            assert isinstance(doc_type, str)
            assert isinstance(patterns, list)
            assert len(patterns) >= 1

    def test_product_aliases_are_strings(self):
        for alias, product in PRODUCT_ALIASES.items():
            assert isinstance(alias, str)
            assert isinstance(product, str)

    def test_product_from_name_is_well_formed(self):
        for product, patterns in PRODUCT_FROM_NAME:
            assert isinstance(product, str)
            assert isinstance(patterns, list)
            assert len(patterns) >= 1


class TestOtcProductDetection:
    """OTCS auto-tagging coverage (Phase 8, Plan 01)."""

    # ── Directory-based detection ──

    def test_webreports_directory(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "webreports" / "report.pdf"
        docs_root.mkdir()
        (docs_root / "webreports").mkdir()
        (docs_root / "webreports" / "report.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "WebReports"

    def test_xecm_directory(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "xecm" / "doc.pdf"
        docs_root.mkdir()
        (docs_root / "xecm").mkdir()
        (docs_root / "xecm" / "doc.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "xECM"

    def test_workflow_directory(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "workflow" / "process.pdf"
        docs_root.mkdir()
        (docs_root / "workflow").mkdir()
        (docs_root / "workflow" / "process.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "Workflow"

    def test_cside_directory(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "cside" / "ide.pdf"
        docs_root.mkdir()
        (docs_root / "cside").mkdir()
        (docs_root / "cside" / "ide.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "CSIDE"

    def test_contentserver_directory(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "contentserver" / "server.pdf"
        docs_root.mkdir()
        (docs_root / "contentserver").mkdir()
        (docs_root / "contentserver" / "server.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "ContentServer"

    def test_brava_directory(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "brava" / "view.pdf"
        docs_root.mkdir()
        (docs_root / "brava").mkdir()
        (docs_root / "brava" / "view.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "Brava"

    def test_ot2_directory(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "ot2" / "ot.pdf"
        docs_root.mkdir()
        (docs_root / "ot2").mkdir()
        (docs_root / "ot2" / "ot.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "OT2"

    # ── Filename-based detection (fallback) ──

    def test_webreports_from_filename(self, tmp_path):
        """'3-0117 Content Server WebReport Design.pdf' → WebReports"""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "3-0117 Content Server WebReport Design.pdf"
        docs_root.mkdir()
        file_path.write_text("")
        assert infer_product(file_path, docs_root) == "WebReports"

    def test_xecm_from_filename(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "xECM-Integration-Guide.pdf"
        docs_root.mkdir()
        file_path.write_text("")
        assert infer_product(file_path, docs_root) == "xECM"

    def test_workflow_from_filename(self, tmp_path):
        docs_root = tmp_path / "docs"
        file_path = docs_root / "Workflow-Admin.pdf"
        docs_root.mkdir()
        file_path.write_text("")
        assert infer_product(file_path, docs_root) == "Workflow"

    # ── Priority tests ──

    def test_directory_takes_priority(self, tmp_path):
        """Directory name wins over filename pattern."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "workflow" / "WebReport-Overview.pdf"
        docs_root.mkdir()
        (docs_root / "workflow").mkdir()
        (docs_root / "workflow" / "WebReport-Overview.pdf").write_text("")
        assert infer_product(file_path, docs_root) == "Workflow"

    def test_existing_products_still_work(self, tmp_path):
        """Regression: AppServer, DataSync, RecordsManagement still resolve."""
        docs_root = tmp_path / "docs"
        docs_root.mkdir()
        # AppServer from directory alias
        p1 = docs_root / "appserver" / "admin.pdf"
        p1.parent.mkdir()
        p1.write_text("")
        assert infer_product(p1, docs_root) == "AppServer"
        # DataSync from filename
        p2 = docs_root / "data-sync-config.pdf"
        p2.write_text("")
        assert infer_product(p2, docs_root) == "DataSync"
        # RecordsManagement from directory
        p3 = docs_root / "records management" / "rm.pdf"
        p3.parent.mkdir()
        p3.write_text("")
        assert infer_product(p3, docs_root) == "RecordsManagement"


class TestInferVendor:
    """Phase 11: Vendor inference tests."""

    def test_vendor_from_filename_opentext(self):
        """'OpenText WebReport Guide.pdf' → 'OpenText' via filename pattern."""
        assert infer_vendor(Path("OpenText WebReport Guide.pdf")) == "OpenText"

    def test_vendor_from_filename_opentext_case_insensitive(self):
        """'OPENTEXT Admin Guide.pdf' → 'OpenText' (case insensitive)."""
        assert infer_vendor(Path("OPENTEXT Admin Guide.pdf")) == "OpenText"

    def test_vendor_from_product_map(self):
        """No vendor in filename, but product='WebReports' → 'OpenText' via VENDOR_MAP."""
        assert (
            infer_vendor(Path("Report Guide.pdf"), "WebReports") == "OpenText"
        )

    def test_vendor_non_opentext_product(self):
        """No patterns matched, product='Adobe' → '' (empty string)."""
        assert infer_vendor(Path("Guide.pdf"), "Adobe") == ""

    def test_vendor_unknown_product(self):
        """No patterns, product='' → '' (empty string)."""
        assert infer_vendor(Path("README.txt"), "") == ""

    def test_vendor_from_ot_prefix(self):
        """'OT-WebReport.pdf' → 'OpenText' via OT- prefix pattern."""
        assert infer_vendor(Path("OT-WebReport.pdf")) == "OpenText"

    def test_vendor_from_ot_underscore_prefix(self):
        """'OT_Admin Guide.pdf' → 'OpenText' via OT_ prefix pattern."""
        assert infer_vendor(Path("OT_Admin Guide.pdf")) == "OpenText"

    def test_vendor_iso_product(self):
        """product='ISO' → 'ISO' via VENDOR_MAP."""
        assert infer_vendor(Path("ISO-9001.pdf"), "ISO") == "ISO"

    def test_vendor_from_directory(self, tmp_path):
        """File in 'OpenText/WebReports/' directory → 'OpenText'."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "OpenText" / "WebReports" / "guide.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_text("")
        assert infer_vendor(file_path) == "OpenText"


class TestInferSubsystem:
    """Phase 11: Subsystem inference tests."""

    def test_subsystem_from_directory(self, tmp_path):
        """File in 'ContentServer/Platform/' → 'Platform' as subsystem."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "ContentServer" / "Platform" / "API.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_text("")
        assert infer_subsystem(file_path, docs_root) == "Platform"

    def test_subsystem_no_intermediate_dir(self, tmp_path):
        """File directly in 'WebReports/file.pdf' → '' (no intermediate dir)."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "WebReports" / "guide.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_text("")
        assert infer_subsystem(file_path, docs_root) == ""

    def test_subsystem_from_filename_pattern(self):
        """'ContentServer API Reference.pdf' → 'API' via SUBSYSTEM_PATTERNS."""
        assert (
            infer_subsystem(Path("ContentServer API Reference.pdf"), Path("/"))
            == "API"
        )

    def test_subsystem_unknown(self):
        """Generic name, no subdirectory → ''."""
        assert infer_subsystem(Path("README.txt"), Path("/")) == ""

    def test_subsystem_skips_varios_dir(self, tmp_path):
        """File in 'varios/ContentServer/misc.pdf' → '' (varios is skipped)."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "varios" / "ContentServer" / "misc.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_text("")
        assert infer_subsystem(file_path, docs_root) == ""

    def test_subsystem_skips_templates_dir(self, tmp_path):
        """File in 'templates/ContentServer/misc.pdf' → '' (templates is skipped)."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "templates" / "ContentServer" / "misc.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_text("")
        assert infer_subsystem(file_path, docs_root) == ""

    def test_subsystem_skips_archive_dir(self, tmp_path):
        """File in 'archive/ContentServer/misc.pdf' → '' (archive is skipped)."""
        docs_root = tmp_path / "docs"
        file_path = docs_root / "archive" / "ContentServer" / "misc.pdf"
        file_path.parent.mkdir(parents=True)
        file_path.write_text("")
        assert infer_subsystem(file_path, docs_root) == ""

    def test_subsystem_security_from_filename(self):
        """Filename with 'security' → 'Security'."""
        assert (
            infer_subsystem(
                Path("ContentServer Security Guide.pdf"), Path("/")
            )
            == "Security"
        )

    def test_subsystem_admin_from_filename(self):
        """Filename with 'admin' → 'Admin' (lower priority but catches)."""
        assert (
            infer_subsystem(Path("Platform Admin Guide.pdf"), Path("/"))
            == "Admin"
        )


class TestExtractDocumentMetadata:
    """Phase 11-02: Document metadata extraction from PDF/DOCX."""

    def test_extract_pdf_metadata_returns_dict(self, tmp_path):
        """Create temp PDF, set metadata, verify extraction returns correct values."""
        import fitz

        doc = fitz.open()
        doc.insert_page(0, text="Test content")
        doc.set_metadata(
            {
                "title": "Test Title",
                "author": "Test Author",
                "subject": "OpenText WebReports",
                "keywords": "admin, guide",
            }
        )
        test_path = tmp_path / "test.pdf"
        doc.save(str(test_path), incremental=False, deflate=True)
        doc.close()

        result = extract_document_metadata(test_path)
        assert result["title"] == "Test Title"
        assert result["author"] == "Test Author"
        assert "OpenText" in result["subject"]
        assert result["keywords"] == "admin, guide"

    def test_extract_docx_metadata_returns_dict(self, tmp_path):
        """Create temp DOCX, set core properties, verify extraction."""
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_paragraph("Test content")
        cp = doc.core_properties
        cp.title = "DOCX Title"
        cp.author = "DOCX Author"
        cp.subject = "OpenText Subject"
        cp.keywords = "docx, test"

        test_path = tmp_path / "test.docx"
        doc.save(str(test_path))

        result = extract_document_metadata(test_path)
        assert result["title"] == "DOCX Title"
        assert result["author"] == "DOCX Author"
        assert "OpenText" in result["subject"]
        assert result["keywords"] == "docx, test"

    def test_extract_unsupported_format(self):
        """.txt file returns empty dict."""
        result = extract_document_metadata(Path("file.txt"))
        assert result == {}

    def test_extract_nonexistent_file(self):
        """Nonexistent file returns empty dict (graceful degradation)."""
        result = extract_document_metadata(Path("/nonexistent/file.pdf"))
        assert result == {}

    def test_build_metadata_text_concatenates_fields(self):
        """All 4 fields concatenated into space-separated string."""
        metadata = {
            "title": "Admin Guide",
            "subject": "OpenText WebReports",
            "author": "Jane Doe",
            "keywords": "admin, guide",
        }
        text = _build_metadata_text(metadata)
        assert "Admin Guide" in text
        assert "OpenText WebReports" in text
        assert "Jane Doe" in text
        assert "admin, guide" in text

    def test_build_metadata_text_with_gaps(self):
        """Only title set returns just the title."""
        metadata = {
            "title": "Only Title",
            "subject": "",
            "author": "",
            "keywords": "",
        }
        text = _build_metadata_text(metadata)
        assert text == "Only Title"

    def test_build_metadata_text_empty(self):
        """Empty dict returns empty string."""
        assert _build_metadata_text({}) == ""


class TestEnrichClassification:
    """Phase 11-02: Gap-filling enrichment from document metadata."""

    def test_enrich_vendor_from_pdf_metadata(self, tmp_path):
        """PDF with subject='OpenText', ambiguous filename → vendor filled to OpenText."""
        import fitz

        doc = fitz.open()
        doc.insert_page(0, text="test")
        doc.set_metadata(
            {
                "title": "General Guide",
                "author": "",
                "subject": "OpenText WebReports",
                "keywords": "",
            }
        )
        pdf_path = tmp_path / "general-guide.pdf"
        doc.save(str(pdf_path), incremental=False, deflate=True)
        doc.close()

        # Initial classification has no vendor (ambiguous filename)
        current = {
            "vendor": "",
            "product": "geral",
            "doc_type": "document",
            "subsystem": "",
        }
        enriched = enrich_classification(pdf_path, tmp_path, current)
        assert enriched["vendor"] == "OpenText"

    def test_enrich_product_from_docx_metadata(self, tmp_path):
        """DOCX with title='WebReport Administrator Guide', generic dir → product filled."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("test")
        cp = doc.core_properties
        cp.title = "WebReport Administrator Guide"
        cp.subject = ""
        cp.author = ""
        cp.keywords = ""

        docx_path = tmp_path / "generic-file.docx"
        doc.save(str(docx_path))

        current = {
            "vendor": "",
            "product": "geral",
            "doc_type": "document",
            "subsystem": "",
        }
        enriched = enrich_classification(docx_path, tmp_path, current)
        assert enriched["product"] == "WebReports"

    def test_enrich_no_override_explicit_classification(self, tmp_path):
        """File in 'WebReports/' dir with clear product → NOT overridden by metadata."""
        import fitz

        doc = fitz.open()
        doc.insert_page(0, text="test")
        doc.set_metadata(
            {
                "title": "Adobe Reader Guide",
                "author": "Adobe Inc",
                "subject": "Adobe",
                "keywords": "",
            }
        )
        webreports_dir = tmp_path / "WebReports"
        webreports_dir.mkdir()
        pdf_path = webreports_dir / "admin-guide.pdf"
        doc.save(str(pdf_path), incremental=False, deflate=True)
        doc.close()

        # Pre-classify to get auto-detected product (WebReports from directory)
        result = classify(pdf_path, tmp_path)
        # Product should remain WebReports, not overridden by Adobe metadata
        assert result["product"] == "WebReports"

    def test_enrich_no_metadata(self, tmp_path):
        """File with no metadata → classification unchanged."""
        current = {
            "vendor": "",
            "product": "geral",
            "doc_type": "document",
            "subsystem": "",
        }
        # Pass a .txt file which has no metadata extraction
        txt_path = tmp_path / "readme.txt"
        txt_path.write_text("hello")
        enriched = enrich_classification(txt_path, tmp_path, current)
        assert enriched == current

    def test_enrich_unsupported_format(self, tmp_path):
        """Unsupported format → no enrichment, unchanged."""
        current = {
            "vendor": "",
            "product": "geral",
            "doc_type": "document",
            "subsystem": "",
        }
        txt_path = tmp_path / "notes.txt"
        txt_path.write_text("notes")
        enriched = enrich_classification(txt_path, tmp_path, current)
        assert enriched == current

    def test_enrich_doc_type_from_metadata(self, tmp_path):
        """PDF subject='Admin Guide', no admin in filename → doc_type enriched."""
        import fitz

        doc = fitz.open()
        doc.insert_page(0, text="test")
        doc.set_metadata(
            {
                "title": "System Reference",
                "author": "",
                "subject": "Admin Guide",
                "keywords": "",
            }
        )
        pdf_path = tmp_path / "reference-doc.pdf"
        doc.save(str(pdf_path), incremental=False, deflate=True)
        doc.close()

        current = {
            "vendor": "",
            "product": "geral",
            "doc_type": "document",
            "subsystem": "",
        }
        enriched = enrich_classification(pdf_path, tmp_path, current)
        assert enriched["doc_type"] == "admin_guide"

    def test_enrich_fills_gaps_when_filename_ambiguous(self, tmp_path):
        """SC2: Ambiguous filename, rich PDF metadata → all gaps filled."""
        import fitz

        doc = fitz.open()
        doc.insert_page(0, text="test")
        doc.set_metadata(
            {
                "title": "WebReport Administrator Guide",
                "author": "",
                "subject": "OpenText Content Server",
                "keywords": "administration, config",
            }
        )
        pdf_path = tmp_path / "doc-12345.pdf"
        doc.save(str(pdf_path), incremental=False, deflate=True)
        doc.close()

        current = {
            "vendor": "",
            "product": "geral",
            "doc_type": "document",
            "subsystem": "",
        }
        enriched = enrich_classification(pdf_path, tmp_path, current)
        assert enriched["vendor"] == "OpenText"
        # Product could be WebReports or ContentServer depending on which
        # metadata field wins — both are valid improvements over "geral"
        assert enriched["product"] != "geral"
        assert enriched["doc_type"] == "admin_guide"
